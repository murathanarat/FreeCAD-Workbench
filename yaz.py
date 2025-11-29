from docx import Document

def kolonun_kayipsiz_kesit_alani(
        path,
        b,
        h,
        L,
        l_,
        i_,
        lambda_,
        f_E_y,
        f_c_0_d,
        C_p,
        sigma_c_0_d
        ):


    doc = Document()
    doc.add_heading("Kolon Hesapları Raporu", level=1)

    # --- Bölüm 1: Temel Hesaplar ---
    doc.add_heading("1. Kesit Özellikleri", level=2)

    table1 = doc.add_table(rows=0, cols=2)
    data1 = [
        ("Kolonun kayıpsız kesit alanı", f"{b}×{h} = {b*h} mm²"),
        ("Kolonun atalet momentleri", f"{(l_)} mm⁴"),
        ("Kolonun atalet yarıçapları", f"{i_} mm"),
        ("Kolonun narinliği", f"{L} / {i_} = {lambda_}"),
        ("Elastik burkulma gerilmesi", f"{f_E_y} MPa"),
        ("Liflere paralel tasarım basınç dayanımı", f"{f_c_0_d} MPa"),
        ("Burkulma katsayısı (c=0.9)", f"{C_p}"),
        ("Kolon burkulma yükü kapasitesi", f"{C_p*f_c_0_d}"),
    ]

    for key, val in data1:
        row = table1.add_row().cells
        row[0].text = key
        row[1].text = val

    # --- Bölüm 2: Yük Altındaki Gerilmeler ---
    doc.add_heading("\n2. Yük Altında Gerilmeler", level=2)

    doc.add_paragraph("232 No’lu eleman (1.0G + 1.0Q + 0.4S + 0.3Ex + Ey + Ez)")
    doc.add_paragraph("49,283 × D = 98.566 kN")
    doc.add_paragraph("83.168 / (140×140) = 5.03 MPa")

    kontrol = doc.add_paragraph()
    kontrol.add_run("Kontrol: 6.83 MPa > 5.03 MPa → ").bold = True
    kontrol.add_run("Kesit Güvenlidir.").italic = True

    # --- Bölüm 3: Yangın Kontrolü ---
    doc.add_heading("\n3. Yangın Kontrolü (30 dk – 3 kenardan etkili)", level=2)

    table2 = doc.add_table(rows=0, cols=2)

    data2 = [
        ("Yangına maruz kalma süresi", "t = 30 dk"),
        ("Kavramsal kömürleşme hızı", "βn = 0.8 mm/dk"),
        ("k0 (Tablo 6.5)", "1.0"),
        ("Kömürleşme derinliği", "dk,n = 24 mm"),
        ("Etkili kömürleşme derinliği", "def = 31 mm"),
        ("Artık kesit çevresi", "P = 0.368 m"),
        ("Artık kesit alanı", "AT = 0.008464 m²"),
        ("Yangın düzeltme katsayısı", "CYN = 0.652"),
        ("Cy20 (Tablo 6.2)", "1.15"),
        ("Etkili kesit alanı", "Ag = 6084 mm²"),
        ("Atalet momenti (etkili)", "3.084×10⁶ mm⁴"),
        ("Atalet yarıçapı (etkili)", "22.51 mm"),
        ("Narinlik", "3000 / 22.51 = 133"),
    ]

    for key, val in data2:
        row = table2.add_row().cells
        row[0].text = key
        row[1].text = val

        

    doc.save(path)
